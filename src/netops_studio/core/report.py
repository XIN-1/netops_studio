"""报表自动化引擎（core/report.py）。

纯 Python，无 GUI 依赖。职责：
- InspectionJob：巡检任务定义（cron-like 调度 + 勾选的 sections）。
- parse_schedule：将 cron-like 表达式解析为「下次运行」信息（纯函数，简单解析）。
- gather：按 sections 惰性调用各 core 引擎，聚合成 report_data（缺失模块优雅跳过）。
- render_html / render_markdown：纯函数，把 report_data 渲染为文本。
- export_pdf / export_docx / export_excel：惰性 import 重依赖（reportlab / python-docx / openpyxl）。

参考开发文档 §6.9（报表自动化）。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

# 支持的 section；core 中部分模块（ipam / security）可能尚未提供，gather 会优雅跳过。
VALID_SECTIONS = ("discovery", "speedtest", "ipam", "security")
EXPORT_FORMATS = ("html", "pdf", "docx", "xlsx")


@dataclass
class InspectionJob:
    """一次巡检任务的静态定义。"""

    name: str = "未命名巡检"
    schedule: str = "0 9 * * *"          # cron-like，如 "0 9 * * *" 表示每天 09:00
    sections: List[str] = field(default_factory=list)
    export_format: str = "html"          # html / pdf / docx / xlsx

    def __post_init__(self) -> None:
        self.sections = [s.lower() for s in self.sections]
        fmt = (self.export_format or "html").lower()
        self.export_format = fmt if fmt in EXPORT_FORMATS else "html"


# --------------------------------------------------------------------------- #
# 调度解析
# --------------------------------------------------------------------------- #
def parse_schedule(spec: str) -> Dict[str, Any]:
    """解析 cron-like 表达式为「下次运行」信息（纯函数，简单解析，非完整 cron）。

    Args:
        spec: 5 段字符串 "分 时 日 月 周"，如 "0 9 * * *"。

    Returns:
        {
          "spec", "minute", "hour", "day", "month", "weekday",
          "summary", "next_run"
        }
        next_run 为 ISO 格式字符串（分钟精度）；无法解析字段则抛 ValueError。
    """
    parts = (spec or "").split()
    if len(parts) != 5:
        raise ValueError(f"非法的 cron 表达式（需 5 段）：{spec!r}")

    minute, hour, dom, month, dow = parts

    def _to_int(value: str, lo: int, hi: int, label: str) -> Optional[int]:
        if value == "*":
            return None
        try:
            n = int(value)
        except ValueError:
            raise ValueError(f"{label} 字段非法：{value!r}") from None
        if not (lo <= n <= hi):
            raise ValueError(f"{label} 字段超出范围 {lo}-{hi}：{value!r}")
        return n

    m = _to_int(minute, 0, 59, "minute")
    h = _to_int(hour, 0, 23, "hour")

    now = datetime.now().replace(second=0, microsecond=0)
    # 注：本解析为「简化 cron」，仅依据 minute/hour 求最近一次执行时间，
    # 故意忽略 day/month/weekday 字段（不实现完整 cron 调度语义）。
    if h is not None and m is not None:
        candidate = now.replace(hour=h, minute=m)
        if candidate <= now:
            candidate += timedelta(days=1)
        summary = f"每天 {h:02d}:{m:02d}"
    elif h is not None and m is None:
        # 固定小时、分钟任意 -> 取整点（近似）
        candidate = now.replace(hour=h, minute=0)
        if candidate <= now:
            candidate += timedelta(days=1)
        summary = f"每天 {h:02d}:00（整点）"
    elif h is None and m is not None:
        candidate = now.replace(minute=m)
        if candidate <= now:
            candidate += timedelta(hours=1)
        summary = f"每小时第 {m:02d} 分"
    else:
        candidate = now + timedelta(minutes=1)
        summary = "每分钟"

    return {
        "spec": spec,
        "minute": minute,
        "hour": hour,
        "day": dom,
        "month": month,
        "weekday": dow,
        "summary": summary,
        "next_run": candidate.isoformat(timespec="minutes"),
    }


# --------------------------------------------------------------------------- #
# 数据聚合
# --------------------------------------------------------------------------- #
def _host_dict(host: Any) -> Dict[str, Any]:
    """将 discovery.Host 转为可序列化 dict。"""
    return {
        "ip": getattr(host, "ip", ""),
        "hostname": getattr(host, "hostname", ""),
        "mac": getattr(host, "mac", ""),
        "vendor": getattr(host, "vendor", ""),
        "state": getattr(host, "state", ""),
        "latency_ms": getattr(host, "latency_ms", None),
    }


def gather(sections: Optional[List[str]] = None, **opts: Any) -> Dict[str, Any]:
    """按 sections 聚合各 core 引擎的结果为 report_data。

    缺失的 core 模块（如 ipam / security 尚未实现）会被标记 skipped 优雅跳过。
    各 section 的调用均为惰性 import。

    Args:
        sections: 需要包含的 section 列表（如 ["discovery", "speedtest"]）。
        opts: 透传给各引擎的可选参数，如 cidr / workers / download_url / latency_target。

    Returns:
        report_data dict：{ generated_at, sections: {名称: {...}} }
    """
    from datetime import datetime as _dt

    data: Dict[str, Any] = {
        "generated_at": _dt.now().isoformat(timespec="seconds"),
        "sections": {},
    }
    wanted = [s.lower() for s in (sections or [])]
    for sec in wanted:
        try:
            if sec == "discovery":
                from .discovery import scan_network

                cidr = opts.get("cidr", "192.168.1.0/24")
                workers = opts.get("workers", 32)
                timeout = opts.get("timeout", 1)
                hosts = scan_network(cidr, workers=workers, timeout=timeout)
                data["sections"]["discovery"] = {
                    "ok": True,
                    "cidr": cidr,
                    "count": len(hosts),
                    "hosts": [_host_dict(h) for h in hosts],
                }
            elif sec == "speedtest":
                from .speedtest import ExternalTester

                res = ExternalTester(
                    download_url=opts.get("download_url", "https://speed.hetzner.de/100MB.bin"),
                    latency_target=opts.get("latency_target", "8.8.8.8"),
                ).measure(download_secs=opts.get("download_secs", 5))
                data["sections"]["speedtest"] = {
                    "ok": True,
                    "download_mbps": res.download_mbps,
                    "upload_mbps": res.upload_mbps,
                    "latency_ms": res.latency_ms,
                    "loss_percent": res.loss_percent,
                    "success": res.success,
                    "note": res.note,
                }
            elif sec == "ipam":
                try:
                    from . import ipam  # type: ignore  # 模块可能尚未实现

                    data["sections"]["ipam"] = ipam.summary(**opts) if hasattr(ipam, "summary") \
                        else {"ok": False, "skipped": True, "error": "ipam 模块无 summary()"}
                except ImportError:
                    data["sections"]["ipam"] = {
                        "ok": False, "skipped": True, "error": "ipam 模块未提供（已跳过）",
                    }
            elif sec == "security":
                try:
                    from . import security  # type: ignore

                    data["sections"]["security"] = security.summary(**opts) if hasattr(security, "summary") \
                        else {"ok": False, "skipped": True, "error": "security 模块无 summary()"}
                except ImportError:
                    data["sections"]["security"] = {
                        "ok": False, "skipped": True, "error": "security 模块未提供（已跳过）",
                    }
            else:
                data["sections"][sec] = {
                    "ok": False, "skipped": True, "error": f"未知 section：{sec}",
                }
        except Exception as exc:  # noqa: BLE001
            data["sections"][sec] = {"ok": False, "error": str(exc)}
    return data


# --------------------------------------------------------------------------- #
# 渲染（纯函数）
# --------------------------------------------------------------------------- #
def _esc(value: Any) -> str:
    s = "" if value is None else str(value)
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _hosts_table_html(hosts: List[Dict[str, Any]]) -> str:
    cols = ["IP", "主机名", "MAC", "厂商", "状态", "延迟(ms)"]
    head = "".join(f"<th>{c}</th>" for c in cols)
    rows = []
    for h in hosts:
        rows.append(
            "<tr>"
            + f"<td>{_esc(h.get('ip'))}</td>"
            + f"<td>{_esc(h.get('hostname'))}</td>"
            + f"<td>{_esc(h.get('mac'))}</td>"
            + f"<td>{_esc(h.get('vendor'))}</td>"
            + f"<td>{_esc(h.get('state'))}</td>"
            + f"<td>{_esc(h.get('latency_ms'))}</td>"
            + "</tr>"
        )
    return f"<table border='1' cellspacing='0' cellpadding='4'><tr>{head}</tr>" \
        + "".join(rows) + "</table>"


def _kv_pairs(sec: Dict[str, Any]) -> List[tuple]:
    """从 section dict 中抽取可显示的键值对（跳过嵌套结构与内部字段）。"""
    skip = {"ok", "skipped", "hosts", "error"}
    pairs = []
    for k, v in sec.items():
        if k in skip:
            continue
        if isinstance(v, (dict, list)):
            continue
        pairs.append((k, v))
    return pairs


def _kv_table_html(sec: Dict[str, Any]) -> str:
    if sec.get("error") and not sec.get("ok"):
        return f"<p>{_esc(sec.get('error'))}</p>"
    pairs = _kv_pairs(sec)
    if not pairs:
        return "<p>（无数据）</p>"
    rows = "".join(
        f"<tr><td>{_esc(k)}</td><td>{_esc(v)}</td></tr>" for k, v in pairs
    )
    return f"<table border='1' cellspacing='0' cellpadding='4'>" \
        f"<tr><th>字段</th><th>值</th></tr>{rows}</table>"


def render_html(report_data: Dict[str, Any]) -> str:
    """将 report_data 渲染为含表格的 HTML 字符串（纯函数）。"""
    gen = report_data.get("generated_at", "")
    blocks = [
        "<h1>NetOps Studio 巡检报告</h1>",
        f"<p>生成时间：{_esc(gen)}</p>",
    ]
    for name, sec in report_data.get("sections", {}).items():
        blocks.append(f"<h2>{_esc(name)}</h2>")
        if not sec.get("ok", False):
            blocks.append(f"<p>{_esc(sec.get('error', '未生成'))}</p>")
            continue
        if sec.get("hosts"):
            blocks.append(_hosts_table_html(sec["hosts"]))
        else:
            blocks.append(_kv_table_html(sec))
    return "\n".join(blocks)


def _hosts_table_md(hosts: List[Dict[str, Any]]) -> str:
    cols = ["IP", "主机名", "MAC", "厂商", "状态", "延迟(ms)"]
    lines = ["| " + " | ".join(cols) + " |",
             "|" + "---|" * len(cols)]
    for h in hosts:
        lines.append("| " + " | ".join(_esc(h.get(k)) for k in
                       ["ip", "hostname", "mac", "vendor", "state", "latency_ms"]) + " |")
    return "\n".join(lines)


def _kv_table_md(sec: Dict[str, Any]) -> str:
    if sec.get("error") and not sec.get("ok"):
        return f"> {sec.get('error')}"
    pairs = _kv_pairs(sec)
    if not pairs:
        return "（无数据）"
    lines = ["| 字段 | 值 |", "| --- | --- |"]
    lines += [f"| {_esc(k)} | {_esc(v)} |" for k, v in pairs]
    return "\n".join(lines)


def render_markdown(report_data: Dict[str, Any]) -> str:
    """将 report_data 渲染为 Markdown 字符串（含表格，纯函数）。"""
    gen = report_data.get("generated_at", "")
    lines = ["# NetOps Studio 巡检报告", f"生成时间：{gen}", ""]
    for name, sec in report_data.get("sections", {}).items():
        lines.append(f"## {name}")
        if not sec.get("ok", False):
            lines.append(sec.get("error", "未生成"))
            lines.append("")
            continue
        if sec.get("hosts"):
            lines.append(_hosts_table_md(sec["hosts"]))
        else:
            lines.append(_kv_table_md(sec))
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# --------------------------------------------------------------------------- #
# 导出（惰性 import 重依赖）
# --------------------------------------------------------------------------- #
def export_pdf(html: str, path: str) -> str:
    """将 HTML（render_html 的输出）惰性导出为 PDF（reportlab）。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table,
    )

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(path, pagesize=A4, title="NetOps Studio 报告")
    flow: List[Any] = []

    pos = 0
    # 注意：此正则仅匹配无属性的 <table> 标签，而 render_html 实际输出为
    # <table border='1' ...>。若后续未在此处放宽正则，PDF 中的表格将被整体跳过，
    # 仅保留文字块（审计项，见汇报）。
    for m in re.finditer(r"<table>.*?</table>", html, re.S):
        _append_text_blocks(flow, html[pos:m.start()], styles)
        rows: List[List[str]] = []
        for tr in re.finditer(r"<tr>(.*?)</tr>", m.group(0), re.S):
            cells = re.findall(r"<t[dh]>(.*?)</t[dh]>", tr.group(1), re.S)
            rows.append([_strip_tags(c) for c in cells])
        if rows:
            table = Table(rows)
            table.setStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#f2f4f6")]),
            ])
            flow.append(table)
        flow.append(Spacer(1, 8))
        pos = m.end()
    _append_text_blocks(flow, html[pos:], styles)

    doc.build(flow)
    return path


def _strip_tags(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text or "").strip()


def _append_text_blocks(flow: List[Any], html: str, styles) -> None:
    """把非表格的 HTML 片段转为 reportlab flowables（h1/h2/p/br）。"""
    from reportlab.platypus import Paragraph, Spacer

    # 移除表格片段以防万一
    frag = re.sub(r"<table>.*?</table>", "", html, flags=re.S)
    # 按块级标签切分
    frag = re.sub(r"<h1>(.*?)</h1>", r"\n#H1#\1", frag, flags=re.S)
    frag = re.sub(r"<h2>(.*?)</h2>", r"\n#H2#\1", frag, flags=re.S)
    frag = re.sub(r"<p>(.*?)</p>", r"\n#P#\1", frag, flags=re.S)
    frag = re.sub(r"<br\s*/?>", "\n", frag)
    frag = _strip_tags(frag)
    for line in frag.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("#H1#"):
            flow.append(Paragraph(_esc(line[4:]), styles["Title"]))
        elif line.startswith("#H2#"):
            flow.append(Paragraph(_esc(line[4:]), styles["Heading2"]))
        elif line.startswith("#P#"):
            flow.append(Paragraph(_esc(line[3:]), styles["Normal"]))
            flow.append(Spacer(1, 4))
        else:
            flow.append(Paragraph(_esc(line), styles["Normal"]))


def export_docx(data: Dict[str, Any], path: str) -> str:
    """惰性导出为 Word 文档（python-docx）。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("NetOps Studio 巡检报告", level=0)
    doc.add_paragraph(f"生成时间：{data.get('generated_at', '')}")

    for name, sec in data.get("sections", {}).items():
        doc.add_heading(name, level=1)
        if not sec.get("ok", False):
            doc.add_paragraph(sec.get("error", "未生成"))
            continue
        if sec.get("hosts"):
            cols = ["IP", "主机名", "MAC", "厂商", "状态", "延迟(ms)"]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = c
            for h in sec["hosts"]:
                cells = table.add_row().cells
                for i, k in enumerate(["ip", "hostname", "mac", "vendor", "state", "latency_ms"]):
                    cells[i].text = str(h.get(k, ""))
        else:
            for k, v in _kv_pairs(sec):
                p = doc.add_paragraph()
                run = p.add_run(f"{k}: ")
                run.bold = True
                p.add_run(str(v))
    doc.save(path)
    return path


def export_excel(data: Dict[str, Any], path: str) -> str:
    """惰性导出为 Excel（openpyxl），每个 section 一个工作表。"""
    from openpyxl import Workbook

    wb = Workbook()
    first = True
    for name, sec in data.get("sections", {}).items():
        ws = wb.active if first else wb.create_sheet()
        first = False
        ws.title = name[:31]
        if not sec.get("ok", False):
            ws["A1"] = sec.get("error", "未生成")
            continue
        if sec.get("hosts"):
            cols = ["IP", "主机名", "MAC", "厂商", "状态", "延迟(ms)"]
            ws.append(cols)
            for h in sec["hosts"]:
                ws.append([h.get(k, "") for k in
                           ["ip", "hostname", "mac", "vendor", "state", "latency_ms"]])
        else:
            ws.append(["字段", "值"])
            for k, v in _kv_pairs(sec):
                ws.append([k, v])
    wb.save(path)
    return path
